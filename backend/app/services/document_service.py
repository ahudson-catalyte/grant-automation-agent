from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, PieChart, Reference
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from icalendar import Calendar, Event, Alarm
from dateutil.parser import parse as date_parser
from datetime import datetime, timedelta
from app.models.schemas import GrantData, Timeline, Budget, WorkPlan
from typing import Dict, Optional
import os
from reportlab.pdfgen import canvas

class DocumentService:
    """Service for generating professional grant management documents"""
    
    def __init__(self, temp_dir: str = "temp_files"):
        self.temp_dir = temp_dir
        os.makedirs(temp_dir, exist_ok=True)
        print(f"DocumentService initialized with temp_dir: {self.temp_dir}")
    
    def generate_workplan_pdf(self, grant_data: GrantData, file_id: str) -> str:
        """Generate comprehensive work plan PDF with timeline and Gantt-style layout"""
        # debug summary
        try:
            print("Generating workplan for:", grant_data.grant_title)
            print("GrantData JSON:", grant_data.json())
        except Exception:
            print("Couldn't print GrantData debug info")
        
        filename = f"{file_id}_workplan.pdf"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc = SimpleDocTemplate(
            filepath, 
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=0.75*inch
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1a365d'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#4a5568'),
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2d3748'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        # Title Page
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("GRANT WORK PLAN", title_style))
        story.append(Spacer(1, 0.1*inch))
        
        grant_title = grant_data.grant_title or "Grant Project"
        story.append(Paragraph(f"<b>{grant_title}</b>", subtitle_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Grant Overview Box
        overview_data = [
            ['GRANT OVERVIEW'],
            ['Organization:', grant_data.organization_name or 'N/A'],
            ['Funder:', grant_data.funder_name or 'N/A'],
            ['Grant Period:', grant_data.grant_period or 'N/A'],
            ['Total Award:', f"${grant_data.grant_amount:,.2f}" if grant_data.grant_amount else 'N/A'],
            ['Document Date:', datetime.now().strftime('%B %d, %Y')],
        ]
        
        overview_table = Table(overview_data, colWidths=[2*inch, 4*inch])
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('SPAN', (0, 0), (-1, 0)),
            ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#edf2f7')),
            ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e0')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        story.append(overview_table)
        story.append(PageBreak())
        
        # Executive Summary
        story.append(Paragraph("EXECUTIVE SUMMARY", heading2_style))
        story.append(Spacer(1, 0.1*inch))
        
        exec_summary = f"""
        This work plan outlines the strategic approach and timeline for implementing the 
        {grant_title}. The plan details specific activities, responsible parties, deliverables, 
        and timelines to ensure successful grant execution and achievement of project objectives.
        """
        story.append(Paragraph(exec_summary, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Work Plan Tasks
        if grant_data.workplan and grant_data.workplan.tasks:
            story.append(Paragraph("DETAILED WORK PLAN", heading2_style))
            story.append(Spacer(1, 0.2*inch))
            
            for idx, task in enumerate(grant_data.workplan.tasks, 1):
                # Task header
                task_header_data = [[f"TASK {idx}: {task.task_name.upper()}"]]
                task_header_table = Table(task_header_data, colWidths=[6.5*inch])
                task_header_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#4299e1')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 11),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 12),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                    ('TOPPADDING', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                ]))
                story.append(task_header_table)
                
                # Task details
                task_details_data = [
                    ['Description:', task.description or 'No description provided'],
                    ['Timeline:', f"{task.start_date or 'TBD'} to {task.end_date or 'TBD'}"],
                    ['Responsible Party:', task.responsible_party or 'To be assigned'],
                    ['Key Deliverables:', task.deliverables or 'See description'],
                ]
                
                task_details_table = Table(task_details_data, colWidths=[1.5*inch, 5*inch])
                task_details_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e2e8f0')),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                    ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e0')),
                ]))
                
                story.append(task_details_table)
                story.append(Spacer(1, 0.3*inch))
        
        # Timeline & Milestones
        if grant_data.timeline and grant_data.timeline.items:
            story.append(PageBreak())
            story.append(Paragraph("KEY MILESTONES & DEADLINES", heading2_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Sort timeline items by date
            sorted_items = sorted(
                grant_data.timeline.items,
                key=lambda x: self._parse_date_safe(x.date)
            )
            
            timeline_data = [['Date', 'Milestone/Deadline', 'Category', 'Amount']]
            
            for item in sorted_items:
                timeline_data.append([
                    item.date,
                    item.description,
                    item.category or 'General',
                    item.amount or '-'
                ])
            
            timeline_table = Table(timeline_data, colWidths=[1.2*inch, 3*inch, 1.3*inch, 1*inch])
            timeline_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ALIGN', (0, 1), (0, -1), 'CENTER'),
                ('ALIGN', (1, 1), (1, -1), 'LEFT'),
                ('ALIGN', (2, 1), (2, -1), 'CENTER'),
                ('ALIGN', (3, 1), (3, -1), 'RIGHT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e0')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            
            story.append(timeline_table)
        
        # Footer with notes
        story.append(PageBreak())
        story.append(Paragraph("IMPLEMENTATION NOTES", heading2_style))
        story.append(Spacer(1, 0.1*inch))
        
        notes = [
            "Regular monitoring and evaluation will be conducted throughout the grant period.",
            "Progress reports will be submitted according to the funder's requirements.",
            "Any significant deviations from this work plan will require prior approval from the funder.",
            "The project team will meet monthly to review progress and address challenges.",
            "Budget expenditures will be tracked against this work plan to ensure alignment."
        ]
        
        for note in notes:
            story.append(Paragraph(f"• {note}", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
        return filepath
    
    def generate_budget_excel(self, grant_data: GrantData, file_id: str) -> str:
        """Generate comprehensive budget and disbursement Excel with charts"""
        filename = f"{file_id}_budget.xlsx"
        filepath = os.path.join(self.temp_dir, filename)
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Budget"
        ws.append(["Category", "Description", "Amount", "Notes"])
        items = getattr(grant_data.budget, "items", []) or []
        for it in items:
            ws.append([getattr(it, "category", ""), getattr(it, "description", ""), getattr(it, "amount", 0), getattr(it, "notes", "")])
        
        # Remove default sheet
        if 'Sheet' in wb.sheetnames:
            wb.remove(wb['Sheet'])
        
        # Style definitions
        header_fill = PatternFill(start_color="1a365d", end_color="1a365d", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=11)
        subheader_fill = PatternFill(start_color="2c5282", end_color="2c5282", fill_type="solid")
        subheader_font = Font(color="FFFFFF", bold=True, size=10)
        total_fill = PatternFill(start_color="4299e1", end_color="4299e1", fill_type="solid")
        total_font = Font(bold=True, size=11, color="FFFFFF")
        currency_format = '"$"#,##0.00'
        border_thin = Border(
            left=Side(style='thin', color='CBD5E0'),
            right=Side(style='thin', color='CBD5E0'),
            top=Side(style='thin', color='CBD5E0'),
            bottom=Side(style='thin', color='CBD5E0')
        )
        
        # ============================================
        # SHEET 1: Budget Summary
        # ============================================
        ws_summary = wb.create_sheet("Budget Summary")
        
        # Title
        ws_summary['A1'] = 'GRANT BUDGET SUMMARY'
        ws_summary['A1'].font = Font(bold=True, size=16, color="1a365d")
        ws_summary.merge_cells('A1:E1')
        
        # Grant info
        ws_summary['A3'] = 'Organization:'
        ws_summary['B3'] = grant_data.organization_name or 'N/A'
        ws_summary['A4'] = 'Grant Title:'
        ws_summary['B4'] = grant_data.grant_title or 'N/A'
        ws_summary['A5'] = 'Grant Period:'
        ws_summary['B5'] = grant_data.grant_period or 'N/A'
        ws_summary['A6'] = 'Total Award:'
        ws_summary['B6'] = grant_data.grant_amount or 0
        ws_summary['B6'].number_format = currency_format
        
        for row in range(3, 7):
            ws_summary[f'A{row}'].font = Font(bold=True)
            ws_summary[f'A{row}'].fill = PatternFill(start_color="edf2f7", end_color="edf2f7", fill_type="solid")
        
        # Budget breakdown headers
        ws_summary['A9'] = 'BUDGET BREAKDOWN'
        ws_summary.merge_cells('A9:E9')
        ws_summary['A9'].font = header_font
        ws_summary['A9'].fill = header_fill
        ws_summary['A9'].alignment = Alignment(horizontal='center', vertical='center')
        
        headers = ['Category', 'Description', 'Amount', 'Percentage', 'Notes']
        for col_num, header in enumerate(headers, 1):
            cell = ws_summary.cell(row=10, column=col_num)
            cell.value = header
            cell.fill = subheader_fill
            cell.font = subheader_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border_thin
        
        # Budget data
        row = 11
        total_amount = 0
        grand_total = grant_data.budget.total_grant_amount if grant_data.budget else 0
        
        if grant_data.budget and grant_data.budget.items:
            for item in grant_data.budget.items:
                ws_summary.cell(row=row, column=1, value=item.category).border = border_thin
                ws_summary.cell(row=row, column=2, value=item.description or '-').border = border_thin
                
                amount_cell = ws_summary.cell(row=row, column=3, value=item.amount)
                amount_cell.number_format = currency_format
                amount_cell.border = border_thin
                
                # Percentage
                if grand_total > 0:
                    pct = (item.amount / grand_total) * 100
                    pct_cell = ws_summary.cell(row=row, column=4, value=pct/100)
                    pct_cell.number_format = '0.0%'
                else:
                    pct_cell = ws_summary.cell(row=row, column=4, value=0)
                pct_cell.border = border_thin
                
                ws_summary.cell(row=row, column=5, value=item.timeline or '-').border = border_thin
                
                total_amount += item.amount
                row += 1
        
        # Total row
        ws_summary.cell(row=row, column=1, value='TOTAL').font = total_font
        ws_summary.cell(row=row, column=1).fill = total_fill
        ws_summary.cell(row=row, column=1).border = border_thin
        
        ws_summary.cell(row=row, column=2, value='').fill = total_fill
        ws_summary.cell(row=row, column=2).border = border_thin
        
        total_cell = ws_summary.cell(row=row, column=3, value=total_amount)
        total_cell.number_format = currency_format
        total_cell.font = total_font
        total_cell.fill = total_fill
        total_cell.border = border_thin
        
        ws_summary.cell(row=row, column=4, value='100%').font = total_font
        ws_summary.cell(row=row, column=4).fill = total_fill
        ws_summary.cell(row=row, column=4).border = border_thin
        
        ws_summary.cell(row=row, column=5, value='').fill = total_fill
        ws_summary.cell(row=row, column=5).border = border_thin
        
        # Column widths
        ws_summary.column_dimensions['A'].width = 20
        ws_summary.column_dimensions['B'].width = 35
        ws_summary.column_dimensions['C'].width = 15
        ws_summary.column_dimensions['D'].width = 12
        ws_summary.column_dimensions['E'].width = 25
        
        # Add pie chart
        if grant_data.budget and grant_data.budget.items and len(grant_data.budget.items) > 0:
            pie = PieChart()
            labels = Reference(ws_summary, min_col=1, min_row=11, max_row=row-1)
            data = Reference(ws_summary, min_col=3, min_row=10, max_row=row-1)
            pie.add_data(data, titles_from_data=True)
            pie.set_categories(labels)
            pie.title = "Budget Distribution"
            pie.height = 10
            pie.width = 15
            ws_summary.add_chart(pie, "G9")
        
        # ============================================
        # SHEET 2: Disbursement Schedule
        # ============================================
        ws_disbursement = wb.create_sheet("Disbursement Schedule")
        
        ws_disbursement['A1'] = 'DISBURSEMENT SCHEDULE'
        ws_disbursement['A1'].font = Font(bold=True, size=16, color="1a365d")
        ws_disbursement.merge_cells('A1:F1')
        
        ws_disbursement['A3'] = 'This schedule tracks all planned and actual grant payments and disbursements.'
        ws_disbursement.merge_cells('A3:F3')
        
        # Headers
        disburse_headers = ['Date', 'Description', 'Planned Amount', 'Actual Amount', 'Status', 'Notes']
        for col_num, header in enumerate(disburse_headers, 1):
            cell = ws_disbursement.cell(row=5, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border_thin
        
        # Add timeline payment items
        row = 6
        if grant_data.timeline and grant_data.timeline.items:
            for item in sorted(grant_data.timeline.items, key=lambda x: self._parse_date_safe(x.date)):
                if item.amount or item.category == 'payment':
                    ws_disbursement.cell(row=row, column=1, value=item.date).border = border_thin
                    ws_disbursement.cell(row=row, column=2, value=item.description).border = border_thin
                    
                    # Planned amount
                    planned_cell = ws_disbursement.cell(row=row, column=3, value=item.amount or 0)
                    planned_cell.number_format = currency_format
                    planned_cell.border = border_thin
                    
                    # Actual amount (empty for tracking)
                    actual_cell = ws_disbursement.cell(row=row, column=4, value=0)
                    actual_cell.number_format = currency_format
                    actual_cell.border = border_thin
                    actual_cell.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
                    
                    # Status dropdown
                    status_cell = ws_disbursement.cell(row=row, column=5, value='Pending')
                    status_cell.border = border_thin
                    status_cell.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
                    
                    # Notes
                    notes_cell = ws_disbursement.cell(row=row, column=6, value='')
                    notes_cell.border = border_thin
                    notes_cell.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
                    
                    row += 1
        
        # Column widths
        ws_disbursement.column_dimensions['A'].width = 15
        ws_disbursement.column_dimensions['B'].width = 40
        ws_disbursement.column_dimensions['C'].width = 15
        ws_disbursement.column_dimensions['D'].width = 15
        ws_disbursement.column_dimensions['E'].width = 15
        ws_disbursement.column_dimensions['F'].width = 30
        
        # ============================================
        # SHEET 3: Expense Tracking
        # ============================================
        ws_expenses = wb.create_sheet("Expense Tracking")
        
        ws_expenses['A1'] = 'EXPENSE TRACKING LOG'
        ws_expenses['A1'].font = Font(bold=True, size=16, color="1a365d")
        ws_expenses.merge_cells('A1:H1')
        
        ws_expenses['A3'] = 'Use this sheet to track individual expenses against budget categories.'
        ws_expenses.merge_cells('A3:H3')
        
        # Headers
        expense_headers = ['Date', 'Vendor/Payee', 'Description', 'Category', 'Amount', 'Payment Method', 'Receipt #', 'Notes']
        for col_num, header in enumerate(expense_headers, 1):
            cell = ws_expenses.cell(row=5, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border_thin
        
        # Add sample row with formulas
        ws_expenses.cell(row=6, column=1, value=datetime.now().strftime('%Y-%m-%d')).border = border_thin
        ws_expenses.cell(row=6, column=2, value='Example Vendor').border = border_thin
        ws_expenses.cell(row=6, column=3, value='Sample expense').border = border_thin
        ws_expenses.cell(row=6, column=4, value='Personnel').border = border_thin
        
        sample_amount = ws_expenses.cell(row=6, column=5, value=0)
        sample_amount.number_format = currency_format
        sample_amount.border = border_thin
        sample_amount.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
        
        for col in range(6, 9):
            cell = ws_expenses.cell(row=6, column=col, value='')
            cell.border = border_thin
            cell.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
        
        # Column widths
        ws_expenses.column_dimensions['A'].width = 12
        ws_expenses.column_dimensions['B'].width = 25
        ws_expenses.column_dimensions['C'].width = 35
        ws_expenses.column_dimensions['D'].width = 20
        ws_expenses.column_dimensions['E'].width = 15
        ws_expenses.column_dimensions['F'].width = 15
        ws_expenses.column_dimensions['G'].width = 15
        ws_expenses.column_dimensions['H'].width = 30
        
        # ============================================
        # SHEET 4: Budget vs Actual
        # ============================================
        ws_variance = wb.create_sheet("Budget vs Actual")
        
        ws_variance['A1'] = 'BUDGET VARIANCE ANALYSIS'
        ws_variance['A1'].font = Font(bold=True, size=16, color="1a365d")
        ws_variance.merge_cells('A1:F1')
        
        # Headers
        variance_headers = ['Category', 'Budgeted', 'Actual Spent', 'Remaining', 'Variance', '% Spent']
        for col_num, header in enumerate(variance_headers, 1):
            cell = ws_variance.cell(row=3, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border_thin
        
        # Add budget categories with tracking formulas
        row = 4
        if grant_data.budget and grant_data.budget.items:
            for item in grant_data.budget.items:
                ws_variance.cell(row=row, column=1, value=item.category).border = border_thin
                
                # Budgeted
                budgeted_cell = ws_variance.cell(row=row, column=2, value=item.amount)
                budgeted_cell.number_format = currency_format
                budgeted_cell.border = border_thin
                
                # Actual (for user to fill)
                actual_cell = ws_variance.cell(row=row, column=3, value=0)
                actual_cell.number_format = currency_format
                actual_cell.border = border_thin
                actual_cell.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
                
                # Remaining formula
                remaining_cell = ws_variance.cell(row=row, column=4, value=f'=B{row}-C{row}')
                remaining_cell.number_format = currency_format
                remaining_cell.border = border_thin
                
                # Variance formula
                variance_cell = ws_variance.cell(row=row, column=5, value=f'=B{row}-C{row}')
                variance_cell.number_format = currency_format
                variance_cell.border = border_thin
                
                # % Spent formula
                pct_cell = ws_variance.cell(row=row, column=6, value=f'=IF(B{row}=0,0,C{row}/B{row})')
                pct_cell.number_format = '0.0%'
                pct_cell.border = border_thin
                
                row += 1
        
        # Total row
        ws_variance.cell(row=row, column=1, value='TOTAL').font = total_font
        ws_variance.cell(row=row, column=1).fill = total_fill
        ws_variance.cell(row=row, column=1).border = border_thin
        
        for col in range(2, 7):
            cell = ws_variance.cell(row=row, column=col)
            if col == 2:
                cell.value = f'=SUM(B4:B{row-1})'
            elif col == 3:
                cell.value = f'=SUM(C4:C{row-1})'
            elif col == 4:
                cell.value = f'=SUM(D4:D{row-1})'
            elif col == 5:
                cell.value = f'=SUM(E4:E{row-1})'
            elif col == 6:
                cell.value = f'=IF(B{row}=0,0,C{row}/B{row})'
                cell.number_format = '0.0%'
            else:
                cell.value = ''
            
            if col in [2, 3, 4, 5]:
                cell.number_format = currency_format
            
            cell.font = total_font
            cell.fill = total_fill
            cell.border = border_thin
        
        # Column widths
        for col in ['A', 'B', 'C', 'D', 'E', 'F']:
            ws_variance.column_dimensions[col].width = 18
        
        wb.save(filepath)
        print(f"✓ Budget excel written: {filepath}")  
        return filepath
    
    def generate_report_template_docx(self, grant_data: GrantData, file_id: str) -> str:
        """Generate comprehensive progress report template in Word with professional formatting"""
        filename = f"{file_id}_report_template.docx"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Helper function to add colored heading
        def add_colored_heading(text, level=1, color='1a365d'):
            heading = doc.add_heading(text, level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.color.rgb = RGBColor(
                    int(color[0:2], 16),
                    int(color[2:4], 16),
                    int(color[4:6], 16)
                )
            return heading
        
        # Title Page
        title = doc.add_heading('Grant Progress Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(26, 54, 93)
        
        doc.add_paragraph()
        
        subtitle = doc.add_paragraph(grant_data.grant_title or '[Grant Title]')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.bold = True
        subtitle_run.font.color.rgb = RGBColor(74, 85, 104)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Grant Information Table
        add_colored_heading('Grant Information', 1)
        
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Organization:', grant_data.organization_name or '[Organization Name]'),
            ('Grant Title:', grant_data.grant_title or '[Grant Title]'),
            ('Funder:', grant_data.funder_name or '[Funder Name]'),
            ('Grant Period:', grant_data.grant_period or '[Grant Period]'),
            ('Reporting Period:', '[Start Date] to [End Date]'),
            ('Report Date:', datetime.now().strftime('%B %d, %Y')),
        ]
        
        for idx, (label, value) in enumerate(info_data):
            info_table.rows[idx].cells[0].text = label
            info_table.rows[idx].cells[1].text = value
            info_table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Add shading to label cells
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'edf2f7')
            info_table.rows[idx].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # Executive Summary
        add_colored_heading('Executive Summary', 1)
        
        exec_para = doc.add_paragraph()
        exec_para.add_run('[Provide a brief overview (2-3 paragraphs) of the grant activities during this reporting period. ')
        exec_para.add_run('Highlight key achievements, challenges encountered, and overall progress toward grant objectives. ')
        exec_para.add_run('This section should give the reader a quick understanding of the project status.]')
        exec_para.runs[0].font.italic = True
        exec_para.runs[0].font.color.rgb = RGBColor(113, 128, 150)
        
        doc.add_paragraph()
        doc.add_paragraph('[Your executive summary here...]')
        doc.add_paragraph()
        
        # Progress Toward Goals
        add_colored_heading('Progress Toward Grant Objectives', 1)
        
        doc.add_paragraph('[For each major objective or goal outlined in your grant proposal, describe the progress made during this reporting period.]')
        doc.add_paragraph()
        
        # Activities and Accomplishments
        add_colored_heading('Activities and Accomplishments', 1)
        
        if grant_data.workplan and grant_data.workplan.tasks:
            doc.add_paragraph('Below is the progress for each planned activity:')
            doc.add_paragraph()
            
            for idx, task in enumerate(grant_data.workplan.tasks, 1):
                # Task heading
                task_heading = add_colored_heading(f'Activity {idx}: {task.task_name}', 2, '2d3748')
                
                # Planned activity
                doc.add_paragraph(f'Planned Activity: {task.description}')
                
                # Progress section (for user to fill)
                progress_heading = doc.add_paragraph('Progress Made:')
                progress_heading.runs[0].font.bold = True
                
                doc.add_paragraph('[Describe specific progress made on this activity during the reporting period]', style='List Bullet')
                doc.add_paragraph('[Include measurable outcomes, participant numbers, or other quantifiable results]', style='List Bullet')
                doc.add_paragraph('[Note any milestones achieved or deliverables completed]', style='List Bullet')
                
                # Challenges section
                challenges_heading = doc.add_paragraph('Challenges Encountered:')
                challenges_heading.runs[0].font.bold = True
                
                doc.add_paragraph('[Describe any obstacles or challenges faced]', style='List Bullet')
                doc.add_paragraph('[Explain how challenges were addressed or plans to address them]', style='List Bullet')
                
                # Next steps
                next_heading = doc.add_paragraph('Next Steps:')
                next_heading.runs[0].font.bold = True
                
                doc.add_paragraph('[Outline planned activities for the next reporting period]', style='List Bullet')
                
                doc.add_paragraph()
        else:
            doc.add_paragraph('Activity 1: [Activity Name]', style='Heading 2')
            doc.add_paragraph('Planned Activity: [Description of planned activity]')
            doc.add_paragraph()
            
            doc.add_paragraph('Progress Made:', style='Heading 3')
            doc.add_paragraph('[Describe progress made during this period]', style='List Bullet')
            doc.add_paragraph('[Include specific accomplishments and outcomes]', style='List Bullet')
            doc.add_paragraph('[Note any challenges or deviations from the plan]', style='List Bullet')
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # Outcomes and Impact
        add_colored_heading('Outcomes and Impact', 1)
        
        impact_para = doc.add_paragraph()
        impact_para.add_run('[Describe the measurable outcomes and impact of your grant activities. ')
        impact_para.add_run('Include both quantitative data (numbers served, items distributed, etc.) and qualitative impact (testimonials, case studies, observed changes).]')
        impact_para.runs[0].font.italic = True
        impact_para.runs[0].font.color.rgb = RGBColor(113, 128, 150)
        
        doc.add_paragraph()
        
        doc.add_paragraph('Quantitative Outcomes:', style='Heading 2')
        doc.add_paragraph('• Number of individuals served: [Number]')
        doc.add_paragraph('• Number of activities/events conducted: [Number]')
        doc.add_paragraph('• Participant satisfaction rate: [Percentage]')
        doc.add_paragraph('• [Other relevant metrics]')
        doc.add_paragraph()
        
        doc.add_paragraph('Qualitative Impact:', style='Heading 2')
        doc.add_paragraph('[Describe the quality of outcomes and changes observed]')
        doc.add_paragraph()
        doc.add_paragraph('Success Stories:', style='Heading 3')
        doc.add_paragraph('[Share 1-2 brief success stories or participant testimonials that illustrate the impact of your work]')
        doc.add_paragraph()
        
        # Financial Report
        doc.add_page_break()
        add_colored_heading('Financial Report', 1)
        
        doc.add_paragraph('This section provides an overview of grant expenditures during the reporting period.')
        doc.add_paragraph()
        
        # Financial summary table
        if grant_data.budget and grant_data.budget.items:
            financial_table = doc.add_table(rows=len(grant_data.budget.items) + 2, cols=5)
            financial_table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = ['Budget Category', 'Total Budgeted', 'Spent to Date', 'Remaining', '% Spent']
            header_cells = financial_table.rows[0].cells
            for idx, header in enumerate(headers):
                header_cells[idx].text = header
                header_cells[idx].paragraphs[0].runs[0].font.bold = True
                
                # Add shading
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '2c5282')
                header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
                
                # White text
                header_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
            # Data rows
            for idx, item in enumerate(grant_data.budget.items, 1):
                row_cells = financial_table.rows[idx].cells
                row_cells[0].text = item.category
                row_cells[1].text = f"${item.amount:,.2f}"
                row_cells[2].text = '$[Amount]'
                row_cells[3].text = '$[Amount]'
                row_cells[4].text = '[%]'
            
            # Total row
            total_cells = financial_table.rows[-1].cells
            total_cells[0].text = 'TOTAL'
            total_cells[0].paragraphs[0].runs[0].font.bold = True
            total_cells[1].text = f"${grant_data.budget.total_grant_amount:,.2f}"
            total_cells[1].paragraphs[0].runs[0].font.bold = True
            total_cells[2].text = '$[Total Spent]'
            total_cells[2].paragraphs[0].runs[0].font.bold = True
            total_cells[3].text = '$[Total Remaining]'
            total_cells[3].paragraphs[0].runs[0].font.bold = True
            total_cells[4].text = '[%]'
            total_cells[4].paragraphs[0].runs[0].font.bold = True
            
            # Shade total row
            for cell in total_cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'e2e8f0')
                cell._element.get_or_add_tcPr().append(shading_elm)
        
        doc.add_paragraph()
        
        doc.add_paragraph('Budget Notes:', style='Heading 2')
        doc.add_paragraph('[Explain any significant variances from the budget (typically anything over 10%)]', style='List Bullet')
        doc.add_paragraph('[Describe any budget reallocations or modifications made during this period]', style='List Bullet')
        doc.add_paragraph('[Note any anticipated budget changes for the next reporting period]', style='List Bullet')
        doc.add_paragraph()
        
        # Challenges and Lessons Learned
        doc.add_page_break()
        add_colored_heading('Challenges and Lessons Learned', 1)
        
        doc.add_paragraph('Challenges:', style='Heading 2')
        doc.add_paragraph('[Describe major challenges encountered during this reporting period]')
        doc.add_paragraph('[Explain how these challenges impacted project activities or outcomes]')
        doc.add_paragraph()
        
        doc.add_paragraph('Solutions and Adaptations:', style='Heading 2')
        doc.add_paragraph('[Describe solutions implemented to address challenges]')
        doc.add_paragraph('[Explain any adaptations made to the project plan]')
        doc.add_paragraph()
        
        doc.add_paragraph('Lessons Learned:', style='Heading 2')
        doc.add_paragraph('[Share key insights gained during this period]')
        doc.add_paragraph('[Describe how these lessons will inform future work]')
        doc.add_paragraph()
        
        # Looking Ahead
        add_colored_heading('Looking Ahead', 1)
        
        doc.add_paragraph('Upcoming Activities:', style='Heading 2')
        doc.add_paragraph('[Outline planned activities for the next reporting period]')
        doc.add_paragraph()
        
        if grant_data.timeline and grant_data.timeline.items:
            doc.add_paragraph('Upcoming Milestones:', style='Heading 2')
            
            # Get future dates
            today = datetime.now()
            future_items = [
                item for item in grant_data.timeline.items
                if self._parse_date_safe(item.date) > today
            ]
            
            if future_items:
                for item in sorted(future_items, key=lambda x: self._parse_date_safe(x.date))[:5]:
                    doc.add_paragraph(f'• {item.date}: {item.description}')
            else:
                doc.add_paragraph('[List upcoming milestones and deadlines]', style='List Bullet')
        
        doc.add_paragraph()
        
        doc.add_paragraph('Anticipated Needs or Support:', style='Heading 2')
        doc.add_paragraph('[Describe any support or resources needed from the funder]')
        doc.add_paragraph('[Note any anticipated challenges or concerns for the next period]')
        doc.add_paragraph()
        
        # Conclusion
        add_colored_heading('Conclusion', 1)
        doc.add_paragraph('[Provide a brief concluding statement that summarizes the overall status of the grant and reaffirms your commitment to achieving the grant objectives.]')
        doc.add_paragraph()
        
        # Attachments
        doc.add_page_break()
        add_colored_heading('Attachments and Supporting Documents', 1)
        
        doc.add_paragraph('The following attachments are included with this report:')
        doc.add_paragraph()
        
        doc.add_paragraph('☐ Financial statements or invoices', style='List Bullet')
        doc.add_paragraph('☐ Photos from events or activities', style='List Bullet')
        doc.add_paragraph('☐ Participant surveys or evaluation data', style='List Bullet')
        doc.add_paragraph('☐ Press coverage or media mentions', style='List Bullet')
        doc.add_paragraph('☐ Letters of support or testimonials', style='List Bullet')
        doc.add_paragraph('☐ Other: [Specify]', style='List Bullet')
        doc.add_paragraph()
        
        # Signature block
        doc.add_paragraph()
        doc.add_paragraph()
        
        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.rows[0].cells[0].text = 'Prepared by:'
        sig_table.rows[0].cells[1].text = 'Date:'
        sig_table.rows[1].cells[0].text = '\n_______________________________'
        sig_table.rows[1].cells[1].text = '\n_______________________________'
        sig_table.rows[2].cells[0].text = '[Name and Title]'
        sig_table.rows[2].cells[1].text = ''
        
        for row in sig_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
        
        doc.save(filepath)
        return filepath
    
    def _parse_date_safe(self, date_str: Optional[str]) -> Optional[datetime]:
        """Return a datetime or None (never raise)."""
        if not date_str:
            return None
        try:
            return date_parser.parse(date_str)
        except Exception:
            try:
                return datetime.strptime(date_str, "%Y-%m-%d")
            except Exception:
                print(f"Warning: unable to parse date: {date_str!r}")
                return None
    
    def generate_calendar_ics(self, grant_data: GrantData, file_id: str) -> str:
        """Generate comprehensive ICS calendar file with all deadlines, meetings, and reminders"""
        filename = f"{file_id}_calendar.ics"
        filepath = os.path.join(self.temp_dir, filename)
        
        cal = Calendar()
        cal.add('prodid', '-//Grant Management Calendar//mxm.dk//')
        cal.add('version', '2.0')
        cal.add('calscale', 'GREGORIAN')
        cal.add('method', 'PUBLISH')
        cal.add('x-wr-calname', f"Grant: {grant_data.grant_title or 'Grant Calendar'}")
        cal.add('x-wr-timezone', 'America/New_York')
        cal.add('x-wr-caldesc', f"Important dates and deadlines for {grant_data.grant_title or 'grant project'}")
        
        events_added = 0
        
        if grant_data.timeline and grant_data.timeline.items:
            for item in grant_data.timeline.items:
                try:
                    event_date = self._parse_date_safe(item.date)
                    
                    # Create event
                    event = Event()
                    
                    # Set summary (title)
                    summary = f"{item.description}"
                    if item.category:
                        summary = f"[{item.category.upper()}] {summary}"
                    event.add('summary', summary)
                    
                    # Set date (all-day event)
                    event.add('dtstart', event_date.date())
                    event.add('dtend', (event_date + timedelta(days=1)).date())
                    
                    # Description
                    description_parts = [item.description]
                    if item.amount:
                        description_parts.append(f"Amount: {item.amount}")
                    if item.category:
                        description_parts.append(f"Category: {item.category}")
                    if grant_data.organization_name:
                        description_parts.append(f"Organization: {grant_data.organization_name}")
                    
                    event.add('description', '\n'.join(description_parts))
                    
                    # Location
                    event.add('location', grant_data.organization_name or '')
                    
                    # Categories/tags
                    categories = ['Grant Management']
                    if item.category:
                        categories.append(item.category.title())
                    event.add('categories', categories)
                    
                    # Priority (higher for payments and compliance)
                    if item.category in ['payment', 'compliance', 'deliverable']:
                        event.add('priority', 1)  # High priority
                    else:
                        event.add('priority', 5)  # Medium priority
                    
                    # Status
                    event.add('status', 'CONFIRMED')
                    
                    # Add multiple alarms/reminders
                    
                    # Reminder 1: Two weeks before
                    alarm1 = Alarm()
                    alarm1.add('action', 'DISPLAY')
                    alarm1.add('description', f"Reminder: {item.description} in 2 weeks")
                    alarm1.add('trigger', timedelta(days=-14))
                    event.add_component(alarm1)
                    
                    # Reminder 2: One week before
                    alarm2 = Alarm()
                    alarm2.add('action', 'DISPLAY')
                    alarm2.add('description', f"Reminder: {item.description} in 1 week")
                    alarm2.add('trigger', timedelta(days=-7))
                    event.add_component(alarm2)
                    
                    # Reminder 3: Three days before (for high priority items)
                    if item.category in ['payment', 'compliance', 'deliverable']:
                        alarm3 = Alarm()
                        alarm3.add('action', 'DISPLAY')
                        alarm3.add('description', f"URGENT: {item.description} in 3 days!")
                        alarm3.add('trigger', timedelta(days=-3))
                        event.add_component(alarm3)
                    
                    # Reminder 4: One day before
                    alarm4 = Alarm()
                    alarm4.add('action', 'DISPLAY')
                    alarm4.add('description', f"Tomorrow: {item.description}")
                    alarm4.add('trigger', timedelta(days=-1))
                    event.add_component(alarm4)
                    
                    # Add UID
                    event.add('uid', f"{file_id}-{events_added}@grantmanagement.local")
                    
                    # Add creation timestamp
                    event.add('dtstamp', datetime.now())
                    
                    cal.add_component(event)
                    events_added += 1
                    
                except Exception as e:
                    print(f"Warning: Could not create calendar event for {item.date}: {e}")
                    continue
        
        # Add recurring monthly review meetings if grant period is known
        if grant_data.grant_period:
            try:
                # Create a recurring event for monthly reviews
                review_event = Event()
                review_event.add('summary', f'Monthly Grant Review - {grant_data.grant_title or "Grant"}')
                review_event.add('description', 
                    'Monthly team meeting to review grant progress, budget, and upcoming milestones.\n\n'
                    'Agenda:\n'
                    '- Review completed activities\n'
                    '- Discuss budget status\n'
                    '- Address challenges\n'
                    '- Plan next month activities'
                )
                
                # Start from first day of next month
                today = datetime.now()
                next_month = today.replace(day=1) + timedelta(days=32)
                review_start = next_month.replace(day=1, hour=10, minute=0, second=0, microsecond=0)
                
                review_event.add('dtstart', review_start)
                review_event.add('dtend', review_start + timedelta(hours=1))
                
                # Make it recurring (monthly)
                review_event.add('rrule', {'freq': 'monthly', 'count': 12})
                
                # Add reminder
                review_alarm = Alarm()
                review_alarm.add('action', 'DISPLAY')
                review_alarm.add('description', 'Grant review meeting in 1 day')
                review_alarm.add('trigger', timedelta(days=-1))
                review_event.add_component(review_alarm)
                
                review_event.add('uid', f"{file_id}-monthly-review@grantmanagement.local")
                review_event.add('dtstamp', datetime.now())
                
                cal.add_component(review_event)
                events_added += 1
                
            except Exception as e:
                print(f"Warning: Could not create monthly review events: {e}")
        
        # Write calendar file
        with open(filepath, 'wb') as f:
            f.write(cal.to_ical())
        
        print(f"✓ Created calendar with {events_added} events")
        return filepath
    
    def _parse_date_safe(self, date_str: str) -> datetime:
        """Safely parse various date formats"""
        from dateutil import parser
        try:
            return parser.parse(date_str)
        except:
            # If parsing fails, return a far future date
            return datetime(2099, 12, 31)
    
    def generate_all_documents(self, grant_data: GrantData, file_id: str, options: Dict[str, bool]) -> Dict[str, str]:
        """Generate all requested documents"""
        generated_files = {}
        
        print(f"\n📄 Generating documents for file_id: {file_id}")
        print(f"Options: {options}")
        
        if options.get('generate_workplan', True):
            try:
                print("  📋 Generating work plan PDF...")
                generated_files['workplan'] = self.generate_workplan_pdf(grant_data, file_id)
                print(f"  ✓ Work plan created: {generated_files['workplan']}")
            except Exception as e:
                print(f"  ❌ Work plan error: {e}")
                generated_files['workplan_error'] = str(e)
        
        if options.get('generate_budget', True):
            try:
                print("  💰 Generating budget Excel...")
                generated_files['budget'] = self.generate_budget_excel(grant_data, file_id)
                print(f"  ✓ Budget created: {generated_files['budget']}")
            except Exception as e:
                print(f"  ❌ Budget error: {e}")
                generated_files['budget_error'] = str(e)
        
        if options.get('generate_report_template', True):
            try:
                print("  📝 Generating report template DOCX...")
                generated_files['report'] = self.generate_report_template_docx(grant_data, file_id)
                print(f"  ✓ Report template created: {generated_files['report']}")
            except Exception as e:
                print(f"  ❌ Report template error: {e}")
                generated_files['report_error'] = str(e)
        
        if options.get('generate_calendar', True):
            try:
                print("  📅 Generating calendar ICS...")
                generated_files['calendar'] = self.generate_calendar_ics(grant_data, file_id)
                print(f"  ✓ Calendar created: {generated_files['calendar']}")
            except Exception as e:
                print(f"  ❌ Calendar error: {e}")
                generated_files['calendar_error'] = str(e)
        
        print(f"✓ Document generation complete. Generated {len([k for k in generated_files.keys() if not k.endswith('_error')])} files\n")
        
        return generated_files